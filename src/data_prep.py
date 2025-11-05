"""
Data preparation utilities for the LLM epistemic belief project.

This script parses the Vesga et al. Study 3 materials, restructures the raw
participant responses, and generates tidy datasets for downstream analysis.

Outputs:
    data/processed/study3_participant_responses.csv
    data/processed/study3_scenarios.csv
"""

from __future__ import annotations

import json
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = ROOT / "data" / "raw"
PROCESSED_DIR = ROOT / "data" / "processed"


# ---------------------------------------------------------------------------
# Utilities for extracting scenario templates from the Word document
# ---------------------------------------------------------------------------

@dataclass
class ScenarioTemplate:
    """Container for a belief framing template."""

    framing: str
    intro_template: str
    statement_a: str
    statement_b: str
    placeholders: Dict[str, str]


def extract_claims_and_actions(doc: Document) -> Dict[str, Dict[str, str]]:
    """Parse the claim/action section of the Study 3 materials."""

    claims: Dict[str, Dict[str, str]] = {}
    current_key: str | None = None
    expected_keys = set(KEY_MAP.values())

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if text.endswith(":") and text[:-1] in {"Procedure", "Stimuli", "Questions", "Control Measure", "Main DV"}:
            # High-level section headers; skip.
            continue

        if ":" in text and not text.startswith(("Truth-Dependent", "Symbolic")):
            key, claim = text.split(":", maxsplit=1)
            key = key.strip()
            if key not in expected_keys:
                continue
            claims[key] = {"claim": claim.strip()}
            current_key = key
            continue

        if text.startswith("Truth-Dependent"):
            assert current_key is not None, "Encountered truth-dependent action without a preceding key."
            claims[current_key]["truth_action"] = text.split(":", maxsplit=1)[1].strip()
            continue

        if text.startswith("Symbolic"):
            assert current_key is not None, "Encountered symbolic action without a preceding key."
            claims[current_key]["symbolic_action"] = text.split(":", maxsplit=1)[1].strip()
            continue

    if not claims:
        raise RuntimeError("Failed to extract claim/action mappings from the Study 3 materials.")

    # Basic validation: ensure each entry has all required fields
    for key, data in claims.items():
        missing = {"claim", "truth_action", "symbolic_action"} - data.keys()
        if missing:
            raise RuntimeError(f"Claim '{key}' is missing fields: {missing}")

    return claims


def extract_topics(doc: Document) -> Dict[str, str]:
    """Extract the topic descriptions for the directionality condition."""

    topics: Dict[str, str] = {}
    capture = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if text.startswith("Topics for the Directionality condition"):
            capture = True
            continue

        if capture:
            if ":" not in text:
                # Stop once we leave the topic list.
                break
            key, topic = text.split(":", maxsplit=1)
            topics[key.strip()] = topic.strip()

    if not topics:
        raise RuntimeError("Failed to extract topics for the directionality condition.")

    return topics


def build_framing_templates(doc: Document) -> Dict[str, ScenarioTemplate]:
    """Create natural-language templates for each framing condition."""

    templates: Dict[str, ScenarioTemplate] = {}

    # Manual extraction guided by document paragraphs. This avoids coupling to
    # exact paragraph indices while keeping the content identical to the stimuli.
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text.startswith("Alex and Blaine were both called") and "[Claim]" in text and "[Topic]" not in text:
            intro_claim = text
            continue

    # After the loop, we may have intro_claim set; if not, raise error
    if "intro_claim" not in locals():
        raise RuntimeError("Unable to locate the base intro paragraph for claim-based conditions.")

    binary_statement = 'Alex: "[Claim]."\nBlaine: "There’s a 98% chance that [Claim]."'
    think_statement = 'Alex: "I believe that [Claim]."\nBlaine: "I think that [Claim]."'

    direction_intro = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Alex and Blaine were both called") and "[Topic]" in text:
            direction_intro = text
            break
    if direction_intro is None:
        raise RuntimeError("Unable to locate the directionality introduction paragraph.")

    direction_statements = (
        'Alex said, "I decided to believe that [Claim]"',
        'Blaine said, "I decided whether to believe that [Claim]"',
    )

    templates["BinaryProbable"] = ScenarioTemplate(
        framing="BinaryProbable",
        intro_template=intro_claim,
        statement_a='Alex: "[Claim]."',
        statement_b='Blaine: "There’s a 98% chance that [Claim]."',
        placeholders={"Claim": "[Claim]"},
    )

    templates["ThinkBelieve"] = ScenarioTemplate(
        framing="ThinkBelieve",
        intro_template=intro_claim,
        statement_a='Alex: "I believe that [Claim]."',
        statement_b='Blaine: "I think that [Claim]."',
        placeholders={"Claim": "[Claim]"},
    )

    templates["Directionality"] = ScenarioTemplate(
        framing="Directionality",
        intro_template=direction_intro,
        statement_a='Alex said, "I decided to believe that [Claim]"',
        statement_b='Blaine said, "I decided whether to believe that [Claim]"',
        placeholders={"Claim": "[Claim]", "Topic": "[Topic]"},
    )

    return templates


# ---------------------------------------------------------------------------
# Data reshaping logic
# ---------------------------------------------------------------------------

VALUE_MAP = {0: -3, 1: -2, 2: -1, 3: 0, 4: 1, 6: 2, 7: 3}

KEY_MAP = {
    "1": "Republican",
    "2": "Jesus",
    "3": "Aliens",
    "4": "Chickens",
    "5": "AI",
}

FRAMING_MAP = {
    "D": "Directionality",
    "TB": "ThinkBelieve",
    "BP": "BinaryProbable",
}

ACTION_MAP = {"t": "Truth Dependent", "s": "Symbolic"}


def tidy_responses(raw_path: Path) -> pd.DataFrame:
    """Convert the wide-format raw CSV into a tidy long dataframe."""

    df = pd.read_csv(raw_path, sep=";")
    records: List[dict] = []

    for column in df.columns:
        if column in {"Part_ID", "Gender", "Age"} or column.startswith("Unnamed"):
            continue

        parts = column.split("_")
        if len(parts) < 4:
            # Skip check columns or unexpected entries
            continue

        stim_key, framing_token, action_token, question = parts[0], parts[1], parts[2], parts[3]

        if framing_token not in FRAMING_MAP or action_token not in ACTION_MAP:
            continue

        framing = FRAMING_MAP[framing_token]
        action_type = ACTION_MAP[action_token]
        question_label = "Action" if question == "Action" else "Certainty" if question == "Certainty" else None

        if question_label is None:
            continue

        key_label = KEY_MAP.get(stim_key)
        if key_label is None:
            continue

        series = df[column].dropna()
        if series.empty:
            continue

        for part_id, raw_value in zip(df.loc[series.index, "Part_ID"], series):
            value = VALUE_MAP.get(raw_value, raw_value)
            records.append(
                {
                    "participant_id": int(part_id),
                    "key": key_label,
                    "framing": framing,
                    "action_type": action_type,
                    "question": question_label,
                    "value": value,
                }
            )

    tidy = pd.DataFrame.from_records(records)
    tidy = tidy.sort_values(["participant_id", "framing", "key", "action_type", "question"]).reset_index(drop=True)
    return tidy


def aggregate_actions(tidy: pd.DataFrame) -> pd.DataFrame:
    """Aggregate action responses per scenario and action type."""

    actions = tidy.loc[tidy["question"] == "Action"].copy()
    grouped = (
        actions.groupby(["framing", "key", "action_type"])
        .agg(
            mean_rating=("value", "mean"),
            median_rating=("value", "median"),
            std_rating=("value", "std"),
            n_responses=("value", "size"),
            share_positive=("value", lambda x: (x > 0).mean()),
            share_negative=("value", lambda x: (x < 0).mean()),
        )
        .reset_index()
    )

    grouped["human_preferred_agent"] = grouped["mean_rating"].apply(
        lambda v: "Blaine" if v > 0 else "Alex" if v < 0 else "Tie"
    )
    grouped["human_preference_strength"] = grouped["mean_rating"].abs()

    return grouped


def assemble_scenarios(
    aggregated: pd.DataFrame,
    templates: Dict[str, ScenarioTemplate],
    claims: Dict[str, Dict[str, str]],
    topics: Dict[str, str],
) -> pd.DataFrame:
    """Attach natural-language context to aggregated statistics."""

    scenario_rows: List[dict] = []

    for _, row in aggregated.iterrows():
        key = row["key"]
        framing = row["framing"]
        action_type = row["action_type"]

        claim_info = claims[key]
        template = templates[framing]

        replacements = {"Claim": claim_info["claim"]}
        if framing == "Directionality":
            replacements["Topic"] = topics[key]

        intro = template.intro_template
        statement_a = template.statement_a
        statement_b = template.statement_b

        for placeholder, source in replacements.items():
            intro = intro.replace(f"[{placeholder}]", source)
            statement_a = statement_a.replace(f"[{placeholder}]", source)
            statement_b = statement_b.replace(f"[{placeholder}]", source)

        action_text = (
            claim_info["truth_action"] if action_type == "Truth Dependent" else claim_info["symbolic_action"]
        )

        scenario_rows.append(
            {
                "key": key,
                "framing": framing,
                "action_type": action_type,
                "claim_text": claim_info["claim"],
                "action_text": action_text,
                "intro_text": intro,
                "alex_statement": statement_a,
                "blaine_statement": statement_b,
                "human_mean_rating": row["mean_rating"],
                "human_median_rating": row["median_rating"],
                "human_std_rating": row["std_rating"],
                "human_preferred_agent": row["human_preferred_agent"],
                "human_preference_strength": row["human_preference_strength"],
                "n_responses": int(row["n_responses"]),
                "share_positive": row["share_positive"],
                "share_negative": row["share_negative"],
            }
        )

    return pd.DataFrame(scenario_rows)


def main() -> None:
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)

    # Load and parse raw assets
    doc = Document(RAW_DIR / "Materials_Study3.docx")
    claims = extract_claims_and_actions(doc)
    topics = extract_topics(doc)
    templates = build_framing_templates(doc)

    tidy = tidy_responses(RAW_DIR / "Check_your_attitudes_Study3_Raw_Wide.csv")
    tidy.to_csv(PROCESSED_DIR / "study3_participant_responses.csv", index=False)

    aggregated = aggregate_actions(tidy)
    scenarios = assemble_scenarios(aggregated, templates, claims, topics)
    scenarios.to_csv(PROCESSED_DIR / "study3_scenarios.csv", index=False)

    # Persist metadata to help analysis scripts.
    metadata = {
        "n_participants": int(tidy["participant_id"].nunique()),
        "n_scenarios": int(scenarios.shape[0] // 2),  # per action type -> divide by 2
        "framing_templates": {k: asdict(v) for k, v in templates.items()},
    }
    (PROCESSED_DIR / "dataset_metadata.json").write_text(json.dumps(metadata, indent=2))


if __name__ == "__main__":
    main()
